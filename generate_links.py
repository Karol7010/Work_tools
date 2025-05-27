from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK

def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param text: The display text for the hyperlink.
    :param url: The URL the hyperlink points to.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    hyperlink.set(qn('w:history'), '1')

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    # Underline
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)

    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def create_doc_with_issues(filename, date_str, issues):
    """
    Create a DOCX with the format you want.

    :param filename: Output filename
    :param date_str: String date for the header (e.g., '2025-05-23')
    :param issues: dict with keys "Blocking issue", "Major issue", "Minor Issue"
                   Values for "Blocking issue" is a list of tuples (list_of_links, optional_text)
                   For other issues, just a list of (text, url) tuples
                   Example:
                   {
                     "Blocking issue": [ ( [("hyperlink1", url1), ("hyperlink2", url2)], " - text"),
                                         ([("hyperlink3", url3)], "") ],
                     "Major issue": [("hyperlink4", url4)],
                     "Minor Issue": [("hyperlink5", url5)]
                   }
    """
    doc = Document()

    # Title line
    doc.add_paragraph(f"Day {date_str} PCC Staging Update")

    for section in ["Blocking issue", "Major issue", "Minor Issue"]:
        if section in issues and issues[section]:
            # Add section header in bold
            p = doc.add_paragraph()
            run = p.add_run(section + ":")
            run.bold = True

            # Add the links list below
            for item in issues[section]:
                # For blocking issue, item can be ([links], extra_text)
                if section == "Blocking issue":
                    links, extra_text = item
                    p = doc.add_paragraph(style='List Bullet')
                    # Add all links separated by commas, then extra_text
                    first_link = True
                    for text, url in links:
                        if not first_link:
                            run_comma = p.add_run(", ")
                        add_hyperlink(p, text, url)
                        first_link = False
                    if extra_text:
                        p.add_run(extra_text)
                else:
                    # For other sections, just a list of (text, url)
                    text, url = item
                    p = doc.add_paragraph(style='List Bullet')
                    add_hyperlink(p, text, url)

    doc.save(filename)

# Example usage:

issues_data = {
    "Blocking issue": [
        ([("hyperlink1", "http://example.com/1"), ("hyperlink2", "http://example.com/2")], " - text"),
        ([("hyperlink3", "http://example.com/3")], "")
    ],
    "Major issue": [
        ("hyperlink4", "http://example.com/4")
    ],
    "Minor Issue": [
        ("hyperlink5", "http://example.com/5")
    ]
}

create_doc_with_issues("output.docx", "2025-05-23", issues_data)
