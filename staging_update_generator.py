"""
describe
"""
# Full path to the source Excel file
source_path = r"C:\Users\eciakar\Ericsson\PCC & PCG & EP5G staging - PCC Minotoring Handovers\PCC staging TR status summary  .xlsm"
staging_csv_file = "PCC staging TR status summary  _2025_PCC_Staging_TRs.csv"

import csv
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime
import os
import shutil
import pandas as pd

def download_csv_file():
    # Specify the exact sheet name you want to export
    sheet_name = "2025_PCC_Staging_TRs"  # <-- Replace this with your actual sheet name

    # Get current folder (where script runs)
    destination_folder = os.getcwd()

    # Extract the Excel filename from source path
    filename = os.path.basename(source_path)

    # Destination full path for the copied Excel file
    destination_path = os.path.join(destination_folder, filename)

    # Copy the Excel file to the script folder
    shutil.copy2(source_path, destination_path)
    print(f"Copied '{filename}' to '{destination_folder}'")

    # Read the specified sheet from the copied Excel file
    df = pd.read_excel(destination_path, sheet_name=sheet_name, engine='openpyxl')

    # Define CSV filename - include the sheet name in the CSV file name for clarity
    csv_filename = f"{os.path.splitext(filename)[0]}_{sheet_name}.csv"
    csv_path = os.path.join(destination_folder, csv_filename)

    # Save as CSV
    df.to_csv(csv_path, index=False)
    print(f"Exported sheet '{sheet_name}' to CSV: '{csv_filename}'")

def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph.
    Compatible with older python-docx versions.
    """

    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Add formatting
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    # Create a w:t element and set the text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    # Append the hyperlink to the paragraph
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def generate_docx(date, blocking, major, minor, filename='update.docx'):
    doc = Document()
    doc.add_heading(f'Day {date} PCC Staging Update', level=1)

    # Blocking issues
    p_title = doc.add_paragraph()
    run = p_title.add_run('Blocking issue:')
    run.bold = True
    for element in blocking:
        p_blocking = doc.add_paragraph(style='ListBullet')
        for job, link in element["links"].items():
            add_hyperlink(p_blocking, job, link)
            p_blocking.add_run(', ')
        # # Remove trailing comma and space
        if p_blocking.runs:
            last_run = p_blocking.runs[-1]
            if last_run.text.endswith(', '):
                last_run.text = last_run.text[:-2]
        p_blocking.add_run(f' - {element["description"]}')

    # Major issue
    p_title = doc.add_paragraph()
    run = p_title.add_run('Major issue:')
    run.bold = True
    for element in major:
        p_major = doc.add_paragraph(style='ListBullet')
        for job, link in element["links"].items():
            add_hyperlink(p_major, job, link)
            p_major.add_run(', ')
        # Remove trailing comma and space
        if p_major.runs:
            last_run = p_major.runs[-1]
            if last_run.text.endswith(', '):
                last_run.text = last_run.text[:-2]
        p_major.add_run(f' - {element["description"]}')

    # Minor issue
    p_title = doc.add_paragraph()
    run = p_title.add_run('Minor issue:')
    run.bold = True
    for element in minor:
        p_minor = doc.add_paragraph(style='ListBullet')
        for job, link in element["links"].items():
            add_hyperlink(p_minor, job, link)
            p_minor.add_run(', ')
        # Remove trailing comma and space
        if p_minor.runs:
            last_run = p_minor.runs[-1]
            if last_run.text.endswith(', '):
                last_run.text = last_run.text[:-2]
        p_minor.add_run(f' - {element["description"]}')
    doc.save(filename)
    print(f"Document saved as {filename}")

def csv_to_issues(row_no):
    with open(staging_csv_file) as csvfile:
        Issues = []
        reader = csv.reader(csvfile)
        next(reader)
        # For master
        for row in reader:
            job = row[row_no].upper().strip()
            headline = row[1]
            comment = row[4]
            if job.strip():
                d = ' '.join([headline, comment])
                d = d.replace('\n', '')
                links = {}
                for single_job in job.split(' '):
                    if "OAM" in single_job:
                        job_no = single_job.replace("OAM", "")
                        links[single_job] = f"https://jenkins-blue-grey.karle005.rnd.gic.ericsson.se/job/staging-tests-jcat-fw--PCC-Staging-OAM-SM/{job_no}/"
                    elif "FT" in single_job:
                        job_no = single_job.replace("FT", "")
                        links[single_job] = f"https://jenkins-blue-grey.karle005.rnd.gic.ericsson.se/job/pcc-staging-deployment--Staging-footprint/{job_no}/"
                    elif "UPG" in single_job:
                        job_no = single_job.replace("UPG", "")
                        links[single_job.strip()] = f"https://jenkins-blue-grey.karle005.rnd.gic.ericsson.se/job/staging-tests-jcat-fw--PCC-Staging-Upgrade-SM/{job_no}/"
                    elif "STAB" in single_job:
                        job_no = single_job.replace("STAB", "")
                        links[single_job] = f"https://jenkins-blue-grey.karle005.rnd.gic.ericsson.se/job/staging-tests-jcat-fw--PCC-Staging-Stability-SM/{job_no}/"
                    else:
                        links[single_job] = "no_link"
                desc_found = False
                for i in range(len(Issues)):
                    if d in Issues[i]:
                        Issues[i]["links"] = Issues[i]["links"] | links
                        desc_found = True
                        break  # if you only want the first match
                if not desc_found:
                    Issues.append({"description": d,"links": links})
        return Issues
if __name__ == '__main__':
    download_csv_file()

    with open(staging_csv_file) as csvfile:
        master_issues = csv_to_issues(13)
        release_issues = csv_to_issues(12)
        all_issues = [*master_issues, *release_issues]
        blocking = []
        major = []
        minor = []
        print(all_issues)

        for issue in all_issues:
            if len(issue['links']) < 3:
                minor.append(issue)
            elif len(issue['links']) < 5:
                major.append(issue)
            else:
                major.append(issue)

        generate_docx(
            date=datetime.date.today(),
            blocking=blocking,
            major=major,
            minor=minor
        )